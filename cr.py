import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import datetime

# =========================
# CONFIGURATION
# =========================

st.set_page_config(
    page_title="CURRICULA-GEN",
    page_icon="🤖",
    layout="wide"
)

st.title("CURRICULA-GEN")

# Use this line on Streamlit Cloud
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

# Use this line locally instead
# GEMINI_API_KEY = "YOUR_API_KEY"

genai.configure(api_key=GEMINI_API_KEY)

model = genai.GenerativeModel("gemini-2.0-flash")

# =========================
# SYSTEM PROMPT (ONE-TIME)
# =========================

SYSTEM_PROMPT = """
You are an expert curriculum designer working in an EdTech environment.

Generate clear, well-structured, and professionally researched learning curricula.

Always follow this structure:
1. Context and Background
2. Course Objectives
3. Module Structure (modules, topics, sub-topics)
4. Learning Outcomes
5. Resources and References (valid functional links only)

Design principles:
- Logical progression from fundamentals to advanced topics
- Alignment between objectives, modules, and outcomes
- Appropriate depth based on course or workshop duration
- Clear and concise explanations

Base the curriculum on industry best practices, reputable educational platforms,
and credible academic or professional resources.
"""

# =========================
# UI — INPUTS
# =========================

st.subheader("Curriculum Parameters")

col1, col2, col3, col4 = st.columns(4)

with col1:
    course_type = st.selectbox(
        "Course Type",
        ["Course", "Workshop"],
        help="Workshop = one-day learning, Course = detailed learning"
    )

with col2:
    num_modules = st.number_input(
        "Number of Modules",
        min_value=1,
        max_value=10,
        value=3
    )

with col3:
    topics_per_module = st.number_input(
        "Topics per Module",
        min_value=1,
        max_value=8,
        value=3
    )

with col4:
    subtopics_per_topic = st.number_input(
        "Max Sub-topics per Topic",
        min_value=0,
        max_value=5,
        value=2
    )

st.subheader("Course Details")

enable_proficiency = st.radio(
    "Select Proficiency Level?",
    ["No", "Yes"],
    horizontal=True
)

proficiency_level = None
if enable_proficiency == "Yes":
    proficiency_level = st.selectbox(
        "Proficiency Level",
        ["Beginner", "Intermediate", "Professional"]
    )

course_topic = st.text_area(
    "Course Topic",
    placeholder="e.g., Introduction to Generative AI",
    height=80
)

primary_resource_url = st.text_input(
    "Primary Resource URL (optional)"
)

# =========================
# PROMPT BUILDER
# =========================

def build_prompt():
    proficiency_text = (
        f"Content should be suitable for a {proficiency_level.lower()} level."
        if proficiency_level else ""
    )

    mode_guidelines = (
        "- Suitable for one-day delivery\n"
        "- Focus on core concepts\n"
        "- Include quick practical activities\n"
        "- Provide time allocation per module"
        if course_type == "Workshop"
        else
        "- Comprehensive topic coverage\n"
        "- Progressive difficulty\n"
        "- Include theoretical and practical components\n"
        "- Emphasize long-term learning outcomes"
    )

    return f"""
{SYSTEM_PROMPT}

Course Type: {course_type}
Course Topic: {course_topic}
Number of Modules: {num_modules}
Topics per Module: {topics_per_module}
Maximum Sub-topics per Topic: {subtopics_per_topic}
Primary Resource URL: {primary_resource_url or "Not provided"}

Mode-specific requirements:
{mode_guidelines}

{proficiency_text}
"""

# =========================
# GENERATION
# =========================

if st.button("Generate Curriculum Structure"):
    if not course_topic.strip():
        st.warning("Please enter a course topic.")
    else:
        with st.spinner("Generating curriculum..."):
            try:
                prompt = build_prompt()
                response = model.generate_content(prompt)

                curriculum_text = response.text
                st.markdown(curriculum_text)

                # =========================
                # WORD DOCUMENT EXPORT
                # =========================

                doc = Document()
                doc.add_heading(f"{course_type}: {course_topic}", 0)
                doc.add_paragraph(
                    f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                )
                doc.add_paragraph(curriculum_text)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="Download as Word Document",
                    data=buffer,
                    file_name=f"curriculum_{course_topic.lower().replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(
                    "⚠️ Gemini API quota or rate limit reached. "
                    "Please wait a minute and try again."
                )

# =========================
# SIDEBAR
# =========================

st.sidebar.markdown("""
## How to Use
1. Enter a course topic
2. Configure modules and topics
3. (Optional) Choose proficiency level
4. Click **Generate Curriculum Structure**
5. Download the curriculum as a Word file

## Output Includes
- Context & background
- Clear objectives
- Structured modules
- Learning outcomes
- Verified learning resources
""")
